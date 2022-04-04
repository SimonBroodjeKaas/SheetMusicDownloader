from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import NoSuchElementException, StaleElementReferenceException
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
import requests
import os
import glob
from docx import Document
from docx.shared import Inches, Cm

driver = None
images = []

def get_element_by_xpath(xpath):
  element = None
  while not element:
    try:
      element = driver.find_element_by_xpath(xpath)
    except NoSuchElementException:
      pass
  
  return element

def get_images():
  global images
  container = driver.find_element_by_id('jmuse-scroller-component')
  divs = container.find_elements_by_xpath(".//*")
  for div in divs:
    try:
      childs = div.find_elements_by_xpath(".//img")
      if len(childs) > 0:
        src = childs[0].get_attribute("src")
        if not src in images:
          images.append(src)
    except StaleElementReferenceException:
      pass
  print(images)

def remove_all_images():
  files = glob.glob('images\\*.jpg')
  for file in files:
    os.remove(file)

def download_image(name, url):
  img_data = requests.get(url).content
  with open(f'{name}.jpg', 'wb') as handler:
      handler.write(img_data)

def download_images(images):
  for i, url in enumerate(images):
    if url != None:
      download_image(f'images/{i+1}', url)
  
def create_word_document():
  os.remove('MuseScore.docx')
  doc = Document()
  p = doc.add_paragraph()
  # r.add_text('Good Morning every body,This is my ')
  # picPath = 'D:/Development/Python/aa.png'
  # r.add_text(' do you like it?')
  sections = doc.sections
  for section in sections:
      section.top_margin = Cm(0)
      section.bottom_margin = Cm(0)
      section.left_margin = Cm(0)
      section.right_margin = Cm(0)
  files = glob.glob('images\\*.jpg')
  for file in files:
    # os.remove(file)
    # r = p.add_run()
    # r.add_picture(file,width=Cm(15), height=Cm(30))
    r = p.add_run()
    r.add_picture(file, width=Cm(21.5), height=Cm(27.5))
  # p = tables[1].rows[0].cells[0].add_paragraph()
  # r = p.add_run()
  # r.add_picture('teste.png',width=Inches(4.0), height=Inches(.7))
  doc.save('MuseScore.docx')
  os.startfile('MuseScore.docx')


def download_sheet_music(url):
  global driver
  global images
  options = Options()
  options.add_argument('--headless')
  driver = webdriver.Chrome(ChromeDriverManager().install(), chrome_options=options)
  driver.set_window_size(1000, 700)

  driver.get(url)

  get_element_by_xpath('/html/body/div[1]/div/div/div/div[2]/div/button[2]').click()

  # target = driver.find_element_by_id('jmuse-scroller-component')
  driver.execute_script('var elem = document.getElementById("jmuse-scroller-component")')


  height = driver.find_element_by_id("jmuse-scroller-component").get_attribute('scrollHeight')

  for scroll in range(0, int(height), int(1200/8)):
    driver.execute_script(f'document.getElementById("jmuse-scroller-component").scrollTop = {scroll}')
    get_images()

  driver.close()

  images = [img for img in images if img != None] #remove None elements

  remove_all_images()

  download_images(images)

  create_word_document()

if __name__ == '__main__':
  download_sheet_music('https://musescore.com/user/96606/scores/152043')